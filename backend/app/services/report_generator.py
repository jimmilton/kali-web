"""Report generation service with PDF and DOCX support."""
from __future__ import annotations

import hashlib
import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from jinja2 import Environment, FileSystemLoader, select_autoescape
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload
from weasyprint import CSS, HTML

from app.models.asset import Asset
from app.models.project import Project
from app.models.report import Report, ReportFormat, ReportTemplate
from app.models.vulnerability import Vulnerability, VulnerabilitySeverity

logger = logging.getLogger(__name__)

# Template directory
TEMPLATE_DIR = Path(__file__).parent.parent / "templates" / "reports"


class ReportGenerator:
    """Generate security assessment reports in various formats."""

    def __init__(self, db: AsyncSession, report: Report):
        self.db = db
        self.report = report
        self.project: Project | None = None
        self.vulnerabilities: list[Vulnerability] = []
        self.assets: list[Asset] = []
        self.env = Environment(
            loader=FileSystemLoader(str(TEMPLATE_DIR)),
            autoescape=select_autoescape(["html", "xml"]),
        )
        self.env.filters["severity_color"] = self._severity_color
        self.env.filters["severity_badge"] = self._severity_badge

    @staticmethod
    def _severity_color(severity: str) -> str:
        """Return color for severity level."""
        colors = {
            "critical": "#dc2626",
            "high": "#ea580c",
            "medium": "#ca8a04",
            "low": "#16a34a",
            "info": "#2563eb",
        }
        return colors.get(severity.lower(), "#6b7280")

    @staticmethod
    def _severity_badge(severity: str) -> str:
        """Return HTML badge for severity level."""
        color = ReportGenerator._severity_color(severity)
        return f'<span style="background-color: {color}; color: white; padding: 2px 8px; border-radius: 4px; font-weight: bold;">{severity.upper()}</span>'

    async def load_data(self) -> None:
        """Load all required data for the report."""
        # Load project
        result = await self.db.execute(
            select(Project).where(Project.id == self.report.project_id)
        )
        self.project = result.scalar_one_or_none()

        if not self.project:
            raise ValueError(f"Project {self.report.project_id} not found")

        # Build vulnerability query
        vuln_query = select(Vulnerability).where(
            Vulnerability.project_id == self.report.project_id
        )

        content = self.report.content or {}

        # Apply filters from content
        if content.get("vulnerability_ids"):
            vuln_query = vuln_query.where(
                Vulnerability.id.in_(content["vulnerability_ids"])
            )

        if content.get("severity_filter"):
            vuln_query = vuln_query.where(
                Vulnerability.severity.in_(content["severity_filter"])
            )

        if content.get("status_filter"):
            vuln_query = vuln_query.where(
                Vulnerability.status.in_(content["status_filter"])
            )

        vuln_query = vuln_query.options(selectinload(Vulnerability.asset))
        vuln_query = vuln_query.order_by(
            # Order by severity (critical first)
            Vulnerability.severity.desc(),
            Vulnerability.created_at.desc(),
        )

        result = await self.db.execute(vuln_query)
        self.vulnerabilities = list(result.scalars().all())

        # Load assets
        asset_query = select(Asset).where(Asset.project_id == self.report.project_id)

        if content.get("asset_ids"):
            asset_query = asset_query.where(Asset.id.in_(content["asset_ids"]))

        asset_query = asset_query.order_by(Asset.type, Asset.value)
        result = await self.db.execute(asset_query)
        self.assets = list(result.scalars().all())

    def _get_context(self) -> dict[str, Any]:
        """Build template context."""
        content = self.report.content or {}
        branding = self.report.branding or {}

        # Count vulnerabilities by severity
        severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
        for vuln in self.vulnerabilities:
            sev = vuln.severity.lower()
            if sev in severity_counts:
                severity_counts[sev] += 1

        # Count assets by type
        asset_counts: dict[str, int] = {}
        for asset in self.assets:
            asset_counts[asset.type] = asset_counts.get(asset.type, 0) + 1

        # Calculate risk score
        risk_score = (
            severity_counts["critical"] * 40
            + severity_counts["high"] * 20
            + severity_counts["medium"] * 10
            + severity_counts["low"] * 5
        )

        return {
            "report": self.report,
            "project": self.project,
            "vulnerabilities": self.vulnerabilities,
            "assets": self.assets,
            "severity_counts": severity_counts,
            "asset_counts": asset_counts,
            "risk_score": min(risk_score, 100),
            "generated_at": datetime.utcnow(),
            "branding": branding,
            "content": content,
            "sections": content.get(
                "sections",
                [
                    "executive_summary",
                    "methodology",
                    "scope",
                    "findings",
                    "recommendations",
                ],
            ),
            # Severity ordering for sorting
            "severity_order": {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4},
        }

    def _get_template_name(self) -> str:
        """Get template filename based on report template type."""
        template_map = {
            ReportTemplate.EXECUTIVE.value: "executive.html",
            ReportTemplate.TECHNICAL.value: "technical.html",
            ReportTemplate.COMPLIANCE.value: "compliance.html",
            ReportTemplate.VULNERABILITY.value: "vulnerability.html",
            ReportTemplate.ASSET.value: "asset.html",
            ReportTemplate.CUSTOM.value: "custom.html",
        }
        return template_map.get(self.report.template, "technical.html")

    async def generate_html(self) -> str:
        """Generate HTML report content."""
        await self.load_data()
        context = self._get_context()
        template_name = self._get_template_name()

        try:
            template = self.env.get_template(template_name)
        except Exception:
            # Fall back to base template
            template = self.env.get_template("base.html")

        return template.render(**context)

    async def generate_pdf(self) -> tuple[bytes, int, str]:
        """Generate PDF report. Returns (content, size, hash)."""
        html_content = await self.generate_html()
        branding = self.report.branding or {}

        # Custom CSS for PDF
        primary_color = branding.get("primary_color", "#1e40af")
        css = CSS(
            string=f"""
            @page {{
                size: A4;
                margin: 2cm;
                @top-center {{
                    content: "{branding.get('header_text', '')}";
                    font-size: 10px;
                    color: #666;
                }}
                @bottom-center {{
                    content: "{branding.get('footer_text', 'Confidential')} - Page " counter(page) " of " counter(pages);
                    font-size: 10px;
                    color: #666;
                }}
            }}
            body {{
                font-family: 'Helvetica', 'Arial', sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #333;
            }}
            h1 {{ color: {primary_color}; border-bottom: 2px solid {primary_color}; padding-bottom: 10px; }}
            h2 {{ color: {primary_color}; margin-top: 20px; }}
            h3 {{ color: #374151; }}
            .severity-critical {{ color: #dc2626; font-weight: bold; }}
            .severity-high {{ color: #ea580c; font-weight: bold; }}
            .severity-medium {{ color: #ca8a04; font-weight: bold; }}
            .severity-low {{ color: #16a34a; font-weight: bold; }}
            .severity-info {{ color: #2563eb; font-weight: bold; }}
            table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: {primary_color}; color: white; }}
            tr:nth-child(even) {{ background-color: #f9fafb; }}
            .finding {{ page-break-inside: avoid; margin-bottom: 20px; padding: 15px; border: 1px solid #e5e7eb; border-radius: 8px; }}
            .evidence {{ background-color: #f3f4f6; padding: 10px; border-radius: 4px; font-family: monospace; font-size: 9pt; white-space: pre-wrap; word-break: break-all; }}
            .badge {{ display: inline-block; padding: 2px 8px; border-radius: 4px; font-weight: bold; color: white; }}
            .badge-critical {{ background-color: #dc2626; }}
            .badge-high {{ background-color: #ea580c; }}
            .badge-medium {{ background-color: #ca8a04; }}
            .badge-low {{ background-color: #16a34a; }}
            .badge-info {{ background-color: #2563eb; }}
        """
        )

        html = HTML(string=html_content)
        pdf_bytes = html.write_pdf(stylesheets=[css])

        file_hash = hashlib.sha256(pdf_bytes).hexdigest()
        return pdf_bytes, len(pdf_bytes), file_hash

    async def generate_docx(self) -> tuple[bytes, int, str]:
        """Generate DOCX report. Returns (content, size, hash)."""
        await self.load_data()
        context = self._get_context()
        branding = self.report.branding or {}

        doc = Document()

        # Set up styles
        styles = doc.styles
        title_style = styles.add_style("ReportTitle", WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(30, 64, 175)

        heading1_style = styles.add_style("ReportHeading1", WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.font.size = Pt(18)
        heading1_style.font.bold = True
        heading1_style.font.color.rgb = RGBColor(30, 64, 175)

        heading2_style = styles.add_style("ReportHeading2", WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.font.size = Pt(14)
        heading2_style.font.bold = True
        heading2_style.font.color.rgb = RGBColor(55, 65, 81)

        # Cover Page
        if branding.get("company_name"):
            p = doc.add_paragraph(branding["company_name"])
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        title = doc.add_paragraph(self.report.title, style="ReportTitle")
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if self.project:
            project_para = doc.add_paragraph(f"Project: {self.project.name}")
            project_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        date_para = doc.add_paragraph(
            f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}"
        )
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_page_break()

        # Table of Contents placeholder
        doc.add_paragraph("Table of Contents", style="ReportHeading1")
        doc.add_paragraph("(Update field to generate TOC)")
        doc.add_page_break()

        sections = context["sections"]

        # Executive Summary
        if "executive_summary" in sections:
            doc.add_paragraph("Executive Summary", style="ReportHeading1")
            summary = f"""
This security assessment report presents the findings from the security testing conducted on {self.project.name if self.project else 'the target environment'}.

During this assessment, a total of {len(self.vulnerabilities)} vulnerabilities were identified:
- Critical: {context['severity_counts']['critical']}
- High: {context['severity_counts']['high']}
- Medium: {context['severity_counts']['medium']}
- Low: {context['severity_counts']['low']}
- Informational: {context['severity_counts']['info']}

Overall Risk Score: {context['risk_score']}/100
"""
            doc.add_paragraph(summary)
            doc.add_page_break()

        # Scope
        if "scope" in sections:
            doc.add_paragraph("Scope", style="ReportHeading1")
            doc.add_paragraph(
                f"A total of {len(self.assets)} assets were included in this assessment:"
            )

            # Assets table
            if self.assets:
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Type"
                hdr_cells[1].text = "Value"
                hdr_cells[2].text = "Status"

                for asset in self.assets[:50]:  # Limit to 50 assets
                    row_cells = table.add_row().cells
                    row_cells[0].text = asset.type
                    row_cells[1].text = asset.value
                    row_cells[2].text = asset.status

            doc.add_page_break()

        # Findings
        if "findings" in sections:
            doc.add_paragraph("Detailed Findings", style="ReportHeading1")

            for i, vuln in enumerate(self.vulnerabilities, 1):
                # Finding header
                doc.add_paragraph(
                    f"Finding {i}: {vuln.title}", style="ReportHeading2"
                )

                # Severity and status
                doc.add_paragraph(f"Severity: {vuln.severity.upper()}")
                doc.add_paragraph(f"Status: {vuln.status}")

                if vuln.cvss_score:
                    doc.add_paragraph(f"CVSS Score: {vuln.cvss_score}")

                if vuln.cve_ids:
                    doc.add_paragraph(f"CVE IDs: {', '.join(vuln.cve_ids)}")

                # Description
                if vuln.description:
                    doc.add_paragraph("Description:", style="ReportHeading2")
                    doc.add_paragraph(vuln.description)

                # Evidence
                content_config = context["content"]
                if content_config.get("include_evidence") and vuln.evidence:
                    doc.add_paragraph("Evidence:", style="ReportHeading2")
                    doc.add_paragraph(vuln.evidence)

                # Remediation
                if content_config.get("include_remediation") and vuln.remediation:
                    doc.add_paragraph("Remediation:", style="ReportHeading2")
                    doc.add_paragraph(vuln.remediation)

                # References
                if content_config.get("include_references") and vuln.references:
                    doc.add_paragraph("References:", style="ReportHeading2")
                    for ref in vuln.references:
                        doc.add_paragraph(f"• {ref}")

                doc.add_paragraph("")  # Spacing

        # Recommendations
        if "recommendations" in sections:
            doc.add_page_break()
            doc.add_paragraph("Recommendations", style="ReportHeading1")

            if context["severity_counts"]["critical"] > 0:
                doc.add_paragraph(
                    "1. IMMEDIATE ACTION REQUIRED: Address all critical severity findings immediately. "
                    "These vulnerabilities pose an imminent risk to the security of the environment."
                )

            if context["severity_counts"]["high"] > 0:
                doc.add_paragraph(
                    "2. HIGH PRIORITY: Remediate high severity findings within the next 7-14 days. "
                    "These issues could lead to significant security breaches if exploited."
                )

            if context["severity_counts"]["medium"] > 0:
                doc.add_paragraph(
                    "3. PLANNED REMEDIATION: Address medium severity findings as part of regular "
                    "security maintenance within the next 30-60 days."
                )

            doc.add_paragraph(
                "4. Implement a regular vulnerability scanning and patch management program."
            )
            doc.add_paragraph(
                "5. Conduct periodic security assessments to identify new vulnerabilities."
            )

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

        file_hash = hashlib.sha256(docx_bytes).hexdigest()
        return docx_bytes, len(docx_bytes), file_hash

    async def generate(self) -> tuple[bytes, int, str, str]:
        """Generate report in the specified format. Returns (content, size, hash, filename)."""
        report_format = self.report.format

        if report_format == ReportFormat.PDF.value:
            content, size, file_hash = await self.generate_pdf()
            filename = f"{self.report.id}.pdf"
        elif report_format == ReportFormat.DOCX.value:
            content, size, file_hash = await self.generate_docx()
            filename = f"{self.report.id}.docx"
        elif report_format == ReportFormat.HTML.value:
            html_content = await self.generate_html()
            content = html_content.encode("utf-8")
            size = len(content)
            file_hash = hashlib.sha256(content).hexdigest()
            filename = f"{self.report.id}.html"
        elif report_format == ReportFormat.MARKDOWN.value:
            # Generate markdown version
            content, size, file_hash = await self._generate_markdown()
            filename = f"{self.report.id}.md"
        elif report_format == ReportFormat.JSON.value:
            content, size, file_hash = await self._generate_json()
            filename = f"{self.report.id}.json"
        else:
            raise ValueError(f"Unsupported format: {report_format}")

        return content, size, file_hash, filename

    async def _generate_markdown(self) -> tuple[bytes, int, str]:
        """Generate Markdown report."""
        await self.load_data()
        context = self._get_context()
        branding = self.report.branding or {}

        lines = []

        # Title
        lines.append(f"# {self.report.title}")
        lines.append("")
        if branding.get("company_name"):
            lines.append(f"**{branding['company_name']}**")
        lines.append(f"*Generated: {datetime.utcnow().strftime('%B %d, %Y')}*")
        lines.append("")

        # Executive Summary
        lines.append("## Executive Summary")
        lines.append("")
        lines.append(f"Total vulnerabilities found: **{len(self.vulnerabilities)}**")
        lines.append("")
        lines.append("| Severity | Count |")
        lines.append("|----------|-------|")
        for sev, count in context["severity_counts"].items():
            lines.append(f"| {sev.capitalize()} | {count} |")
        lines.append("")
        lines.append(f"**Risk Score:** {context['risk_score']}/100")
        lines.append("")

        # Assets
        lines.append("## Scope")
        lines.append("")
        lines.append(f"Total assets: **{len(self.assets)}**")
        lines.append("")
        if self.assets:
            lines.append("| Type | Value | Status |")
            lines.append("|------|-------|--------|")
            for asset in self.assets[:30]:
                lines.append(f"| {asset.type} | {asset.value} | {asset.status} |")
        lines.append("")

        # Findings
        lines.append("## Detailed Findings")
        lines.append("")
        for i, vuln in enumerate(self.vulnerabilities, 1):
            lines.append(f"### Finding {i}: {vuln.title}")
            lines.append("")
            lines.append(f"**Severity:** {vuln.severity.upper()}")
            lines.append(f"**Status:** {vuln.status}")
            if vuln.cvss_score:
                lines.append(f"**CVSS Score:** {vuln.cvss_score}")
            lines.append("")
            if vuln.description:
                lines.append("#### Description")
                lines.append(vuln.description)
                lines.append("")
            if vuln.remediation:
                lines.append("#### Remediation")
                lines.append(vuln.remediation)
                lines.append("")
            lines.append("---")
            lines.append("")

        content = "\n".join(lines).encode("utf-8")
        file_hash = hashlib.sha256(content).hexdigest()
        return content, len(content), file_hash

    async def _generate_json(self) -> tuple[bytes, int, str]:
        """Generate JSON report."""
        import json

        await self.load_data()

        data = {
            "report": {
                "id": str(self.report.id),
                "title": self.report.title,
                "description": self.report.description,
                "template": self.report.template,
                "format": self.report.format,
                "generated_at": datetime.utcnow().isoformat(),
            },
            "project": {
                "id": str(self.project.id) if self.project else None,
                "name": self.project.name if self.project else None,
                "description": self.project.description if self.project else None,
            },
            "summary": {
                "total_vulnerabilities": len(self.vulnerabilities),
                "total_assets": len(self.assets),
                "severity_counts": {
                    sev: sum(1 for v in self.vulnerabilities if v.severity == sev)
                    for sev in ["critical", "high", "medium", "low", "info"]
                },
            },
            "vulnerabilities": [
                {
                    "id": str(v.id),
                    "title": v.title,
                    "description": v.description,
                    "severity": v.severity,
                    "status": v.status,
                    "cvss_score": float(v.cvss_score) if v.cvss_score else None,
                    "cvss_vector": v.cvss_vector,
                    "cve_ids": v.cve_ids,
                    "cwe_ids": v.cwe_ids,
                    "evidence": v.evidence,
                    "remediation": v.remediation,
                    "references": v.references,
                    "asset_id": str(v.asset_id) if v.asset_id else None,
                }
                for v in self.vulnerabilities
            ],
            "assets": [
                {
                    "id": str(a.id),
                    "type": a.type,
                    "value": a.value,
                    "status": a.status,
                    "metadata": a.metadata_,
                    "tags": a.tags,
                }
                for a in self.assets
            ],
        }

        content = json.dumps(data, indent=2).encode("utf-8")
        file_hash = hashlib.sha256(content).hexdigest()
        return content, len(content), file_hash
